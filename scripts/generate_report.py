"""
FinOps Engagement Report Generator
====================================
Generates a formatted .docx report from a previous engagement's docx as a style
template — inheriting margins, fonts, heading colours, and numbering definitions
without carrying over any of the previous content.

Usage
-----
1. Copy this file into the customer engagement folder (e.g. FinOps/<customer>/).
2. Set TEMPLATE_PATH to an existing engagement docx (any previous report works —
   it is used only for styles, not content).
3. Set OUTPUT_PATH for the new report file.
4. Replace the placeholder sections below with the actual findings.
5. Run:  python generate_report.py

Requirements
------------
    pip install python-docx

Template docx spec (confirmed from previous engagements)
---------------------------------------------------------
- Page:     A4 (8.278" × 11.694")
- Margins:  Top 1.069"  Bottom 0.194"  Left 0.431"  Right 0.340"
- Fonts:    Aptos (body)  |  Aptos Display (H1 20pt bold, H3 12pt bold)
- Heading colours: 1F3864 (H1, H3)
- Tables:   'Table Grid' style, D9D9D9 header fill, bold header text
- Bullets:  List Paragraph style + w:numPr referencing numId=15 in template
            (numId=15 → abstractNumId=2 → bullet character \\uf0b7)
            IMPORTANT: bullets require the numPr XML injection below —
            applying 'List Paragraph' style alone does NOT render bullets.
- Savings:  Bold Normal paragraph: "Total monthly saving £XXX"
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Configuration — set these before running ─────────────────────────────────

TEMPLATE_PATH = r"C:\path\to\FinOps\<previous-customer>\v1.0 Azure Cost Saving - <Customer>.docx"
OUTPUT_PATH   = r"C:\path\to\FinOps\<customer>\v1.0 Azure Cost Saving - <Customer> <Mon>'<YY>.docx"

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set table cell background shading colour."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)


def add_table(doc, headers, rows, col_widths=None):
    """
    Add a Table Grid table with a D9D9D9 header row.

    headers    : list of column header strings
    rows       : list of row tuples/lists — values are coerced to str
    col_widths : optional list of column widths in inches (same length as headers)
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    hdr_row = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = text
        set_cell_bg(cell, 'D9D9D9')
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if para.runs:
            para.runs[0].bold = True

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            row.cells[c_idx].text = str(val) if val is not None else ''

    if col_widths:
        for row in table.rows:
            for c_idx, width in enumerate(col_widths):
                if c_idx < len(row.cells):
                    row.cells[c_idx].width = Inches(width)

    return table


def add_bullet(doc, text):
    """
    Add a properly formatted bullet point (List Paragraph style + numPr XML).

    NOTE: Setting style='List Paragraph' alone is not sufficient — Word requires
    the w:numPr element referencing a numbering definition to render the bullet
    character and indentation. numId=15 references the standard bullet list
    (abstractNumId=2, bullet char \\uf0b7) carried across from the template docx.
    """
    p = doc.add_paragraph(style='List Paragraph')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    numId_el = OxmlElement('w:numId')
    numId_el.set(qn('w:val'), '15')
    numPr.append(ilvl)
    numPr.append(numId_el)
    pPr.append(numPr)
    p.add_run(text)


def add_saving(doc, amount_str):
    """Add a bold 'Total monthly saving £XXX' closing line for a finding."""
    p = doc.add_paragraph()
    p.add_run(f'Total monthly saving \u00a3{amount_str}').bold = True


def add_para(doc, text, bold=False, italic=False, size=None):
    """Add a Normal paragraph with optional bold/italic/size overrides."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    return p


# ── Initialise document from template ────────────────────────────────────────

doc = Document(TEMPLATE_PATH)

# Strip all existing content — keeps styles, margins, numbering definitions
body = doc.element.body
to_remove = [c for c in body if c.tag.split('}')[-1] != 'sectPr']
for el in to_remove:
    body.remove(el)

# ── Report content — replace everything below with actual findings ────────────

doc.add_heading('Report', level=1)

# Overview
doc.add_heading('Overview', level=2)
doc.add_paragraph()
add_para(doc, (
    'The purpose of this report is to show opportunities for cost optimisation identified '
    'during a FinOps engagement for <Customer>.'
))
doc.add_paragraph()
add_para(doc, 'The scope of this report covers the following areas:')
doc.add_paragraph()
add_bullet(doc, '<Finding area 1>')
add_bullet(doc, '<Finding area 2>')
doc.add_paragraph()

# Finding 1
doc.add_heading('<Finding Category>', level=2)
doc.add_heading('<Sub-finding>', level=3)
doc.add_paragraph()
add_para(doc, '<Analysis and recommendation text.>')
doc.add_paragraph()
add_table(doc,
    headers=['Column 1', 'Column 2', 'Column 3'],
    rows=[
        ['Value', 'Value', 'Value'],
    ],
    col_widths=[2.5, 2.5, 2.5]
)
doc.add_paragraph()
add_saving(doc, 'XXX')
doc.add_paragraph()

# Summary
doc.add_heading('Summary', level=2)
add_para(doc, (
    'The following table summarises the potential savings that could be made by implementing '
    'the recommendations in this report.'
))
doc.add_paragraph()
add_table(doc,
    headers=['Finding', 'Monthly Saving', 'Effort', 'Risk'],
    rows=[
        ['<Finding 1>', '\u00a3XXX', 'Low/Medium/High', 'Low/Medium/High'],
        ['<Finding 2>', '\u00a3XXX', 'Low/Medium/High', 'Low/Medium/High'],
    ],
    col_widths=[3.5, 1.5, 1.0, 2.3]
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Total confirmed monthly saving: \u00a3XXX').bold = True

# ── Save ──────────────────────────────────────────────────────────────────────

doc.save(OUTPUT_PATH)
print(f'Saved: {OUTPUT_PATH}')
